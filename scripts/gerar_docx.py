#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera o edital verticalizado em DOCX imprimível/anotável (caixas ☐ por item).

Uso: python3 gerar_docx.py edital.json --out edital.docx
"""
import argparse, json, os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_LOGO_DEFAULT = Path(__file__).resolve().parent.parent / "assets" / "logo.png"
LOGO = os.environ.get("EDITAL_LOGO") or str(_LOGO_DEFAULT)  # opcional: marca própria do usuário
AZUL = RGBColor(0x25, 0x63, 0xEB)
SLATE = RGBColor(0x1E, 0x29, 0x3B)
MUT = RGBColor(0x64, 0x74, 0x8B)


def _page_number(paragraph):
    run = paragraph.add_run()
    for t, val in (("begin", None), ("instr", "PAGE"), ("end", None)):
        if t == "instr":
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = " PAGE "
        else:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), t)
        run._r.append(el)


def build(data, out):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    for m in ("top", "bottom", "left", "right"):
        setattr(sec, f"{m}_margin", Cm(2))

    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
    h1 = doc.styles["Heading 1"]; h1.font.name = "Calibri"; h1.font.size = Pt(15); h1.font.color.rgb = AZUL

    # Cabeçalho
    try:
        doc.add_picture(LOGO, width=Cm(1.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Edital Verticalizado"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = SLATE
    sub = " · ".join(x for x in [data.get("concurso"), data.get("cargo")] if x)
    if sub:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(sub); rr.font.size = Pt(12); rr.font.color.rgb = AZUL
    meta = " · ".join(x for x in [data.get("orgao"), data.get("data_edital")] if x)
    if meta:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(meta); rr.font.size = Pt(9.5); rr.font.color.rgb = MUT

    # Quadro resumo
    doc.add_paragraph()
    qh = doc.add_paragraph(); rq = qh.add_run("Quadro-resumo das disciplinas"); rq.bold = True
    rq.font.size = Pt(13); rq.font.color.rgb = SLATE
    for i, d in enumerate(data.get("disciplinas", []), 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{d.get('nome','')} ").bold = True
        p.add_run(f"({len(d.get('itens', []))} itens)").font.color.rgb = MUT
    total = sum(len(d.get("itens", [])) for d in data.get("disciplinas", []))
    tot = doc.add_paragraph(); rt = tot.add_run(f"Total: {total} itens para acompanhar.")
    rt.italic = True; rt.font.color.rgb = MUT

    # Disciplinas
    for d in data.get("disciplinas", []):
        doc.add_page_break()
        doc.add_heading(d.get("nome", ""), level=1)
        for it in d.get("itens", []):
            nivel = int(it.get("nivel", 1))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3 + (nivel - 1) * 0.7)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("☐  ")  # ☐
            rn = p.add_run(f"{it.get('numero','')}  "); rn.font.color.rgb = MUT
            rn.font.size = Pt(10)
            rt = p.add_run(it.get("texto", ""))
            if nivel == 1:
                rt.bold = True; rt.font.color.rgb = SLATE
            else:
                rt.font.size = Pt(10.5)

    # Rodapé com número de página
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Edital Verticalizado · Simião Cavalcante   |   ")
    fr.font.size = Pt(8); fr.font.color.rgb = MUT
    _page_number(fp)

    doc.save(out)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("json"); ap.add_argument("--out", "-o", required=True)
    a = ap.parse_args()
    data = json.loads(Path(a.json).read_text(encoding="utf-8"))
    build(data, a.out)
    n = sum(len(d["itens"]) for d in data.get("disciplinas", []))
    print(f"OK -> {a.out} | {len(data.get('disciplinas', []))} disciplinas, {n} itens")


if __name__ == "__main__":
    main()
